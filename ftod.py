#!/usr/bin/env python3 
""" 
folder-to-doc.py 
================ 
 
Walk a folder, read every file EXCEPT Word documents (.doc / .docx), and write 
their full contents into a single Word (.docx) file. 
 
The idea: point this at the code folder that is actually deployed, and it 
produces one document containing all of the current code. You can diff that 
document against a previous run to see what changed, or hand it off for review. 
 
USAGE 
----- 
1. Install the one dependency (once): 
 
       pip install python-docx 
 
2. Either edit the INPUT_PATH / OUTPUT_PATH constants below, or pass them on 
   the command line (command line wins): 
 
       python folder-to-doc.py "C:\\path\\to\\deployment\\folder" 
       python folder-to-doc.py "C:\\path\\to\\src" "C:\\path\\to\\out\\Code.docx" 
 
WHAT IT INCLUDES / SKIPS 
------------------------ 
- Skips Word docs: .doc and .docx (as requested). 
- Skips common noise folders (node_modules, .git, __pycache__, dist, build, ...). 
- Skips files that are not readable text (images, binaries) automatically. 
- Skips the output .docx itself if it happens to live inside the input folder. 
Adjust SKIP_DIRS / SKIP_EXTENSIONS below to taste. 
""" 
 
import os 
import sys 
from datetime import datetime 
 
try: 
    from docx import Document 
    from docx.shared import Pt 
    from docx.enum.text import WD_BREAK 
except ImportError: 
    sys.exit( 
        "The 'python-docx' package is required.\n" 
        "Install it with:  pip install python-docx" 
    ) 
 
# --------------------------------------------------------------------------- 
# CONFIG  --  edit these, or override them from the command line. 
# --------------------------------------------------------------------------- 
INPUT_PATH = r"C:\Users\Aditya kumar jha\OneDrive\Documents\Continental" 
OUTPUT_PATH = r"C:\Users\Aditya kumar jha\OneDrive\Documents\Continental\DeploymentCode.docx" 
 
# File formats that are themselves documents -> never read/embed these. 
DOC_EXTENSIONS = {".doc", ".docx"} 
 
# Folders that are never worth dumping into the document. 
SKIP_DIRS = { 
    ".git", "node_modules", "__pycache__", ".vscode", ".idea", 
    "dist", "build", ".next", "coverage", "venv", ".venv", 
} 
 
# Extra file extensions to skip (binaries / assets). Text files that slip 
# through here are still auto-skipped if they don't decode as UTF-8. 
SKIP_EXTENSIONS = { 
    ".png", ".jpg", ".jpeg", ".gif", ".ico", ".svg", ".webp", ".bmp", 
    ".pdf", ".zip", ".gz", ".tar", ".7z", ".rar", 
    ".mp3", ".mp4", ".mov", ".avi", ".woff", ".woff2", ".ttf", ".eot", 
    ".exe", ".dll", ".bin", ".lock", 
} 
 
def should_skip_file(rel_path: str, abs_path: str, output_abspath: str) -> str | None: 
    """Return a reason string if the file should be skipped, else None.""" 
    ext = os.path.splitext(abs_path)[1].lower() 
    if os.path.abspath(abs_path) == output_abspath: 
        return "output document" 
    if ext in DOC_EXTENSIONS: 
        return "Word document (.doc/.docx)" 
    if ext in SKIP_EXTENSIONS: 
        return f"binary/asset ({ext})" 
    return None 
 
def read_text(abs_path: str) -> str | None: 
    """Read a file as UTF-8 text; return None if it isn't decodable text.""" 
    try: 
        with open(abs_path, "r", encoding="utf-8") as fh: 
            return fh.read() 
    except (UnicodeDecodeError, PermissionError, OSError): 
        return None 
 
def collect_files(root: str, output_abspath: str): 
    """Yield (relative_path, absolute_path) for every file to embed.""" 
    for dirpath, dirnames, filenames in os.walk(root): 
        # Prune skip-dirs in place so os.walk doesn't descend into them. 
        dirnames[:] = [d for d in dirnames if d not in SKIP_DIRS] 
        for name in sorted(filenames): 
            abs_path = os.path.join(dirpath, name) 
            rel_path = os.path.relpath(abs_path, root) 
            yield rel_path, abs_path, output_abspath 
 
def add_code_block(doc: "Document", text: str) -> None: 
    """Add monospaced, line-preserving text as a single paragraph.""" 
    if text is None: 
        text = "" 
    # Normalise line endings. 
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n") 
    para = doc.add_paragraph() 
    for i, line in enumerate(lines): 
        run = para.add_run(line) 
        run.font.name = "Consolas" 
        run.font.size = Pt(8) 
        if i < len(lines) - 1: 
            run.add_break(WD_BREAK.LINE) 
 
def build_document(input_path: str, output_path: str) -> None: 
    input_path = os.path.abspath(input_path) 
    output_abspath = os.path.abspath(output_path) 
 
    if not os.path.isdir(input_path): 
        sys.exit(f"Input path is not a folder: {input_path}") 
 
    included = [] 
    skipped = [] 
 
    for rel_path, abs_path, out_abs in collect_files(input_path, output_abspath): 
        reason = should_skip_file(rel_path, abs_path, out_abs) 
        if reason: 
            skipped.append((rel_path, reason)) 
            continue 
        content = read_text(abs_path) 
        if content is None: 
            skipped.append((rel_path, "not UTF-8 text")) 
            continue 
        included.append((rel_path, content)) 
 
    # ---- Build the .docx ------------------------------------------------- 
    doc = Document() 
 
    doc.add_heading("Deployment Code Snapshot", level=0) 
    subtitle = doc.add_paragraph( 
        f"Source: {input_path}\n" 
        f"Generated: {datetime.now():%Y-%m-%d %H:%M}\n" 
        f"{len(included)} files included, {len(skipped)} skipped" 
    ) 
    subtitle.runs[0].italic = True 
 
    # Summary of included files. 
    doc.add_heading("Included files", level=1) 
    for rel_path, _ in included: 
        doc.add_paragraph(rel_path, style="List Bullet") 
 
    # Summary of skipped files (so nothing is silently dropped). 
    if skipped: 
        doc.add_heading("Skipped files", level=1) 
        for rel_path, reason in skipped: 
            doc.add_paragraph(f"{rel_path}  ({reason})", style="List Bullet") 
 
    doc.add_page_break() 
 
    # Per-file full content. 
    for rel_path, content in included: 
        doc.add_heading(rel_path, level=1) 
        doc.add_heading("Full file content", level=2) 
        add_code_block(doc, content) 
        doc.add_page_break() 
 
    os.makedirs(os.path.dirname(output_abspath) or ".", exist_ok=True) 
    doc.save(output_abspath) 
 
    print(f"Created: {output_abspath}") 
    print(f"Included {len(included)} files, skipped {len(skipped)}.") 
 
def main() -> None: 
    input_path = sys.argv[1] if len(sys.argv) > 1 else INPUT_PATH 
    output_path = sys.argv[2] if len(sys.argv) > 2 else OUTPUT_PATH 
    build_document(input_path, output_path) 
 
if __name__ == "__main__": 
    main()